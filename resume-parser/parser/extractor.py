"""
Phase 1: Extract raw text from uploaded resume.
Uses Docling for PDF (layout-aware, handles
two-column, tables, complex formatting).
Uses python-docx for DOCX.
Falls back to plain text for TXT files.
"""

import io
import tempfile
import os
from typing import Tuple


def extract_text_from_bytes(
    file_bytes: bytes,
    filename: str,
    mime_type: str = "",
) -> Tuple[str, str]:
    """
    Extract raw text from file bytes.
    Returns (raw_text, parser_used).
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf") or "pdf" in mime_type:
        return _extract_pdf(file_bytes, filename)
    elif (
        filename_lower.endswith(".docx")
        or "word" in mime_type
        or "officedocument" in mime_type
    ):
        return _extract_docx(file_bytes)
    elif filename_lower.endswith(".doc"):
        return _extract_docx(file_bytes)
    else:
        try:
            text = file_bytes.decode("utf-8", errors="replace")
        except Exception:
            text = file_bytes.decode("latin-1", errors="replace")
        return text, "plain-text"


def _extract_pdf(file_bytes: bytes, filename: str) -> Tuple[str, str]:
    """
    Use Docling for PDF extraction.
    """
    try:
        from docling.document_converter import DocumentConverter

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            converter = DocumentConverter()
            result = converter.convert(tmp_path)
            text = result.document.export_to_markdown()

            import re

            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            text = re.sub(r"\*([^*]+)\*", r"\1", text)
            text = re.sub(r"#{1,6}\s*", "", text)
            text = re.sub(r"\n{3,}", "\n\n", text)

            return text.strip(), "docling"
        finally:
            os.unlink(tmp_path)

    except Exception as e:
        print(f"[extractor] Docling failed: {e}")
        try:
            import fitz

            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages = []
            for page in doc:
                pages.append(page.get_text())
            doc.close()
            return "\n\n".join(pages), "pymupdf"
        except Exception as e2:
            print(f"[extractor] PyMuPDF fallback failed: {e2}")
            return "", "failed"


def _extract_docx(file_bytes: bytes) -> Tuple[str, str]:
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs), "python-docx"
    except Exception as e:
        print(f"[extractor] DOCX extraction failed: {e}")
        return "", "failed"
